from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "Bilingual_Packet_June_2_2026.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="DADCE0", size="4"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        el = borders.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def set_table_width(table, widths):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = width
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width.inches * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def style_doc(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for attr in ("top_margin", "right_margin", "bottom_margin", "left_margin"):
        setattr(section, attr, Inches(0.75))
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, color in [
        ("Heading 1", 17, "1F4D78"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 11, "434343"),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(4)


def add_title(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run("Bilingual Learning Packet / Trousse bilingue")
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0, 0, 0)
    p2 = doc.add_paragraph("Tuesday, June 2, 2026 | Mardi 2 juin 2026")
    p2.paragraph_format.space_after = Pt(10)


def add_small_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(85, 85, 85)


def add_checklist_table(doc):
    doc.add_heading("First Page Checklist / Liste de contrôle", level=1)
    rows = [
        ("Daily quote / Citation du jour", "5 min"),
        ("Chores / Tâches", "35-45 min"),
        ("Oliver study page / Page d'étude d'Oliver", "35-45 min"),
        ("Logan study page / Page d'étude de Logan", "35-45 min"),
        ("Read + answer / Lire + répondre", "15 min"),
        ("Art challenge / Défi artistique", "30-40 min"),
        ("Brain activation game / Jeu pour activer le cerveau", "10-15 min"),
        ("Outside PE or science / Bouger ou explorer dehors", "25-35 min"),
        ("French word + restful activity / Mot français + repos", "15-20 min"),
        ("YouTube research + movie idea / Recherche + film", "35-50 min"),
        ("Journal prompt / Question de journal", "10 min"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    headers = ["Done", "Topic / Sujet", "Time / Temps"]
    for idx, h in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = h
        set_cell_shading(cell, "E8EEF5")
    for topic, time in rows:
        cells = table.add_row().cells
        cells[0].text = "[ ]"
        cells[1].text = topic
        cells[2].text = time
    set_table_width(table, [Inches(0.65), Inches(5.0), Inches(1.35)])
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
    add_small_note(doc, "Total target: about 3 to 4 hours, including chores and breaks. / Objectif : environ 3 à 4 heures, avec les tâches et les pauses.")


def add_daily_quote_and_chores(doc):
    doc.add_heading("Daily Quote / Citation du jour", level=1)
    quote = '"You do not have to see the whole staircase, just take the first step." - Martin Luther King Jr.'
    fr = '"Tu n\'as pas besoin de voir tout l\'escalier; fais seulement le premier pas." - Martin Luther King Jr.'
    p = doc.add_paragraph()
    p.add_run("English: ").bold = True
    p.add_run(quote)
    p = doc.add_paragraph()
    p.add_run("Français : ").bold = True
    p.add_run(fr)

    add_small_note(doc, "Quote selected from the local /daily-quotes list for June 2, 2026.")

    doc.add_heading("Chores / Tâches de la maison", level=1)
    chores = [
        ("Mow lawn", "Tondre la pelouse", "Wear shoes, clear sticks first, ask before starting the mower."),
        ("Sweep and mop kitchen", "Balayer et laver le plancher de la cuisine", "Chairs up, corners first, mop last."),
        ("Make bed", "Faire son lit", "Pillow, blanket, floor clear."),
        ("Laundry", "Lessive", "Sort, start or fold one load."),
        ("Dishes", "Vaisselle", "Load, unload, or hand-wash one sink group."),
    ]
    for en, fr_text, tip in chores:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run("[ ] ").bold = True
        p.add_run(f"{en} / {fr_text}: ")
        p.add_run(tip)


def add_study_page(doc, kid, focus, questions):
    doc.add_page_break()
    french_name = f"d'{kid}" if kid[0].lower() in "aeiouh" else f"de {kid}"
    doc.add_heading(f"{kid}'s Study Moment / Moment d'étude {french_name}", level=1)
    add_small_note(doc, f"Theme: {focus}. Answer in English, French, or both. / Thème : {focus}. Réponds en anglais, en français, ou les deux.")
    for subject, prompt in questions:
        doc.add_heading(f"{subject} / {translate_subject(subject)}", level=2)
        doc.add_paragraph(prompt)
        for _ in range(2):
            line = doc.add_paragraph("_" * 78)
            line.paragraph_format.space_after = Pt(2)


def translate_subject(subject):
    return {
        "Social Studies": "Sciences humaines",
        "Science": "Sciences",
        "Health": "Santé",
        "Math": "Mathématiques",
        "Language Arts": "Français/Anglais",
    }[subject]


def add_reading(doc):
    doc.add_page_break()
    doc.add_heading("Something to Read / Petit texte à lire", level=1)
    text = (
        "At sunrise, Maya found tiny tracks in the garden mud. They curved around the lettuce, "
        "paused by a shiny stone, then disappeared under the fence. She did not know the animal, "
        "so she drew the tracks, measured one print with a twig, and wrote three clues. Later, "
        "her brother checked a field guide. Together they learned that asking careful questions "
        "can turn an ordinary morning into a discovery."
    )
    doc.add_paragraph(text)
    add_small_note(doc, "French helper: les traces = tracks; le jardin = garden; une découverte = a discovery.")
    qs = [
        "What three things did Maya do before checking the field guide? / Quelles trois choses Maya a-t-elle faites?",
        "What is the main lesson of the story? / Quelle est la grande leçon de l'histoire?",
    ]
    for q in qs:
        p = doc.add_paragraph(style="List Number")
        p.add_run(q)
        doc.add_paragraph("_" * 78)


def add_creative_and_active(doc):
    doc.add_heading("Art Challenge / Défi artistique", level=1)
    doc.add_paragraph(
        "Build or create 'A Map of a Brave First Step.' Choose sketch, paint, Blender, LEGO, or Minecraft. Include: a starting place, one obstacle, one helpful tool, and a finish marker. Add labels in English and French."
    )
    doc.add_paragraph(
        "Crée « La carte d'un premier pas courageux ». Choisis dessin, peinture, Blender, LEGO ou Minecraft. Ajoute : le départ, un obstacle, un outil utile et l'arrivée."
    )

    doc.add_heading("Brain Activation Game / Jeu d'activation", level=1)
    doc.add_paragraph(
        "Alphabet Switch: Pick a category such as animals, foods, places, or Minecraft blocks. Take turns naming one item from A to Z. Twist: every third answer must be in French or include a French color."
    )

    doc.add_heading("Outside PE or Science Discovery / Activité dehors", level=1)
    doc.add_paragraph(
        "Ten-Minute Nature Sprint: jog or fast-walk for 3 minutes, then find five textures outside: rough, smooth, wet/damp, bendy, and sharp-looking. Do not touch unsafe items. Sketch or list them, then explain which one would change fastest in rain."
    )


def add_word_rest_research_movie_journal(doc):
    doc.add_page_break()
    doc.add_heading("French Word / Mot français", level=1)
    doc.add_paragraph("curieux / curieuse = curious")
    doc.add_paragraph("Sentence: Je suis curieux parce que je veux comprendre. / I am curious because I want to understand.")

    doc.add_heading("Restful Activity / Activité reposante", level=1)
    doc.add_paragraph(
        "Cloud Breathing: lie down or sit outside. Breathe in for 4, hold for 2, breathe out for 6. After five rounds, name one sound, one color, and one thing your body is thankful for."
    )

    doc.add_heading("YouTube Rabbit-Hole Research / Recherche style terrier", level=1)
    doc.add_paragraph(
        "Topic: How do beavers change a landscape? / Comment les castors changent-ils un paysage?"
    )
    steps = [
        "Search with an adult: 'beaver dam ecosystem time lapse' and 'why beavers are ecosystem engineers'.",
        "Watch 2 short videos, 5-10 minutes each. Pause when you hear a new fact.",
        "Write 5 facts, 2 questions, and 1 thing you would build in Minecraft or LEGO to explain it.",
        "Bonus French words: le castor = beaver; un barrage = dam; l'eau = water.",
    ]
    for step in steps:
        doc.add_paragraph(step, style="List Number")

    doc.add_heading("Movie Suggestion / Idée de film", level=1)
    doc.add_paragraph(
        "Orion and the Dark on Netflix. It is a family movie about facing fears, and Netflix lists English and French audio/subtitle options. Watch in English with French subtitles, or switch one scene to French and catch three familiar words. Source checked June 2, 2026: netflix.com/title/81476885."
    )

    doc.add_heading("Journal Prompt / Question de journal", level=1)
    doc.add_paragraph(
        "What was one 'first step' you took today? What made it hard, what helped, and what would your future self say about it? / Quel premier pas as-tu fait aujourd'hui? Qu'est-ce qui était difficile, qu'est-ce qui a aidé, et que dirait ton futur toi?"
    )


def main():
    doc = Document()
    style_doc(doc)
    add_title(doc)
    add_checklist_table(doc)
    add_daily_quote_and_chores(doc)

    add_study_page(
        doc,
        "Oliver",
        "Communities, energy, choices, patterns, and persuasive words",
        [
            ("Social Studies", "A community has a library, a rink, and a food bank. Which one helps people learn, which one helps people connect, and which one helps people in an emergency? Explain your thinking. / Explique ton raisonnement."),
            ("Science", "Name three forms of energy you used today. For each one, say where it came from and how you noticed it."),
            ("Health", "Make a mini plan for a hard moment: What is one body signal, one calming action, and one person you can ask for help?"),
            ("Math", "A chore playlist is 36 minutes. You spend 1/3 mowing prep, 1/4 sweeping, and the rest folding laundry. How many minutes are left for laundry? Show your work."),
            ("Language Arts", "Write four persuasive sentences convincing someone to try your art challenge. Include one strong verb and one reason."),
        ],
    )

    add_study_page(
        doc,
        "Logan",
        "Maps, habitats, wellness, fractions, and story clues",
        [
            ("Social Studies", "Draw a simple map of the kitchen and yard. Add a compass rose, three labels, and a safe route for carrying laundry or dishes."),
            ("Science", "Choose one backyard habitat: grass, soil, tree, or garden. What living things might use it for food, water, shelter, or space?"),
            ("Health", "Create a 'reset menu' with three healthy choices: one movement, one drink/snack, and one quiet activity."),
            ("Math", "You fold 24 clothing items. 1/2 are shirts, 1/4 are socks, and the rest are towels. How many towels are there? Show your equation."),
            ("Language Arts", "The sentence is: 'The tracks vanished under the fence.' What can you infer? Write two possible explanations and circle your best one."),
        ],
    )

    add_reading(doc)
    add_creative_and_active(doc)
    add_word_rest_research_movie_journal(doc)

    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.text = "Bilingual Learning Packet | Trousse bilingue"
        for run in footer.runs:
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(85, 85, 85)

    doc.save(OUT)


if __name__ == "__main__":
    main()
